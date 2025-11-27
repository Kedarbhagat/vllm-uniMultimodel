import docx
from docx import Document as DocxDocument
import os
import logging
from typing import List, Dict, Any
from langchain_core.documents import Document
import hashlib
import time

from pdfloader import chunk_text, convert_to_langchain_docs

# Add to your existing imports
try:
    import mammoth
    MAMMOTH_AVAILABLE = True
except ImportError:
    MAMMOTH_AVAILABLE = False

logger = logging.getLogger("doc_processor")

def extract_text_from_docx(docx_path: str) -> List[Dict[str, Any]]:
    """Extract text from a Word document (.docx) file"""
    start_time = time.time()
    extracted_chunks = []

    try:
        # Open the document
        doc = DocxDocument(docx_path)

        # Extract text from paragraphs with section tracking
        current_section = "main"
        section_content = []

        # Process paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check if this is likely a heading (simplified approach)
            if para.style.name.startswith('Heading'):
                # If we have accumulated content, save it as a chunk
                if section_content:
                    extracted_chunks.append({
                        "text": "\n".join(section_content),
                        "section": current_section,
                        "source": "docx_direct"
                    })

                # Start a new section
                current_section = text
                section_content = []
            else:
                section_content.append(text)

        # Add the last section if not empty
        if section_content:
            extracted_chunks.append({
                "text": "\n".join(section_content),
                "section": current_section,
                "source": "docx_direct"
            })

        # Process tables (simplified)
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    table_text.append(" | ".join(row_text))

            if table_text:
                extracted_chunks.append({
                    "text": "\n".join(table_text),
                    "section": "table",
                    "source": "docx_table"
                })

        processing_time = time.time() - start_time
        logger.info(f"Processed DOCX in {processing_time:.2f}s: {docx_path}")
        return extracted_chunks

    except Exception as e:
        logger.error(f"Failed to process DOCX {docx_path}: {str(e)}")
        return []

def extract_text_from_doc(doc_path: str) -> List[Dict[str, Any]]:
    """Extract text from legacy Word document (.doc) file using mammoth"""
    if not MAMMOTH_AVAILABLE:
        logger.error("Mammoth library not available. Cannot process .doc files.")
        return []

    start_time = time.time()

    try:
        # Use mammoth to convert .doc to HTML then extract text
        with open(doc_path, "rb") as docfile:
            result = mammoth.convert_to_html(docfile)
            html = result.value

            # Simple HTML processing (in production you'd use BeautifulSoup)
            # Split by common heading tags
            sections = []
            current_section = {"title": "main", "content": []}

            # Very basic HTML parsing - in production use proper HTML parsing
            for line in html.split("<"):
                if line.startswith("h1") or line.startswith("h2") or line.startswith("h3"):
                    # Extract heading text (simplified)
                    heading = line.split(">")[1].split("<")[0].strip()

                    # Save previous section if not empty
                    if current_section["content"]:
                        sections.append(current_section)

                    # Start new section
                    current_section = {"title": heading, "content": []}
                elif line.startswith("p") or line.startswith("div"):
                    # Extract paragraph text (simplified)
                    content = line.split(">")[1].split("<")[0].strip()
                    if content:
                        current_section["content"].append(content)

            # Add final section
            if current_section["content"]:
                sections.append(current_section)

            # Convert sections to our standard format
            extracted_chunks = []
            for section in sections:
                if section["content"]:
                    extracted_chunks.append({
                        "text": "\n".join(section["content"]),
                        "section": section["title"],
                        "source": "doc_mammoth"
                    })

        processing_time = time.time() - start_time
        logger.info(f"Processed DOC in {processing_time:.2f}s: {doc_path}")
        return extracted_chunks

    except Exception as e:
        logger.error(f"Failed to process DOC {doc_path}: {str(e)}")
        return []

def convert_word_to_langchain_docs(
    doc_path: str,
    max_sentences: int = 10
) -> List[Document]:
    """Process Word documents (.doc or .docx) to LangChain Documents"""
    start_time = time.time()
    logger.info(f"Starting processing of Word document: {doc_path}")

    # Determine file type
    filename = os.path.basename(doc_path)
    file_ext = os.path.splitext(filename)[1].lower()

    try:
        # Choose appropriate extraction method
        if file_ext == '.docx':
            extracted = extract_text_from_docx(doc_path)
        elif file_ext == '.doc':
            extracted = extract_text_from_doc(doc_path)
        else:
            logger.error(f"Unsupported file type: {file_ext}")
            return []

        documents = []

        # Create document chunks
        for entry in extracted:
            # Use the same chunking function as for PDFs
            chunks = chunk_text(entry["text"], max_sentences=max_sentences)

            for chunk_idx, chunk in enumerate(chunks):
                if not chunk.strip():
                    continue

                # Create unique ID for the chunk
                chunk_id = f"{filename}_{entry.get('section', 'main')}_c{chunk_idx}"

                # Create LangChain Document
                doc = Document(
                    page_content=chunk,
                    metadata={
                        "source": doc_path,
                        "filename": filename,
                        "section": entry.get("section", "main"),
                        "chunk_id": chunk_id,
                        "extraction_method": entry["source"],
                        "processing_time": time.time() - start_time
                    }
                )
                documents.append(doc)

        logger.info(f"Completed processing Word document {doc_path}: {len(documents)} chunks created")
        return documents

    except Exception as e:
        logger.error(f"Failed to process Word document {doc_path}: {str(e)}")
        return []

# In wordloader.py, update the process_document function:
def process_document(file_path: str, **kwargs) -> List[Document]:
    """Process any supported document type (PDF, DOCX, DOC, XLSX, XLS)"""
    file_ext = os.path.splitext(file_path)[1].lower()

    try:
        if file_ext in ['.pdf']:
            return convert_to_langchain_docs(file_path, **kwargs)
        elif file_ext in ['.docx', '.doc']:
            return convert_word_to_langchain_docs(file_path, **kwargs)
        elif file_ext in ['.xlsx', '.xls']:
            # Import Excel loader when needed (avoids circular import)
            try:
                from excel_loader import convert_excel_to_langchain_docs
                logger.info(f"Successfully imported excel_loader for {file_path}")
                result = convert_excel_to_langchain_docs(file_path, **kwargs)
                logger.info(f"Excel processing returned {len(result)} documents")
                return result
            except ImportError as e:
                logger.error(f"Failed to import excel_loader: {e}")
                raise
            except Exception as e:
                logger.error(f"Excel processing failed: {e}")
                raise
        else:
            logger.error(f"Unsupported file extension: {file_ext}")
            return []
    except Exception as e:
        logger.error(f"Document processing failed for {file_path}: {str(e)}")
        raise